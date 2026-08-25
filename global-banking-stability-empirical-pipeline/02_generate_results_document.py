"""
02_generate_results_document.py
================================
PROJECT: Global Banking Sector Stability Around Systemic Crises (2000-2023)

PURPOSE OF THIS SCRIPT:
Takes the clean panel produced by 01_fetch_data.py and runs the FULL
empirical pipeline on it, writing every result as a TABLE into a single
Word document. No figures/plots are created here on purpose - this is the
"raw results" document. A later script (03_generate_figures.py) will turn
selected results into publication-style charts.

WHAT IT DOES, STEP BY STEP:

  TABLE 1 - Descriptive statistics
      Mean/median/std/min/max of every variable, split by income group
      (High income vs Emerging market). First look at the data.

  TABLE 2 - Structural break test (Chow Test)
      Tests whether the *global average* Bank Z-score series has a
      statistically significant break in 2008 (GFC) and separately in
      2020 (COVID). Formal way of asking "did the trend actually change?"

  TABLE 3 - Panel Fixed Effects vs Random Effects + Hausman test
      Regresses Bank Z-score on capital ratio, NPL ratio, ROA, liquidity,
      GDP growth and inflation, with country and year fixed effects.
      Runs both FE and RE, then the Hausman test tells us which one is
      the econometrically correct choice for this data.

  TABLE 4 - Difference-in-Differences (2008 GFC)
      Treatment = Emerging market economies, Control = High income
      economies. Tests whether emerging markets' banking stability fell
      MORE than high-income countries' did after 2008, i.e. whether
      financial development level acted as a buffer.

  TABLE 5 - Difference-in-Differences (2020 COVID)
      Same design, applied to the 2020 shock, so the two crises can be
      compared side by side.

  TABLE 6 - Event study (relative-time table)
      Average Z-score in each income group for years -3 to +3 relative to
      each crisis - lets you see the pattern year by year, not just
      before/after.

  TABLE 7 - Robustness checks
      - Clustered standard errors by country (already used above, shown
        again for transparency)
      - Placebo test: pretend the "crisis" happened in a year with no
        actual shock (2004) - if DiD still finds a significant effect,
        that would be a red flag; finding NO effect supports the main
        result.
      - Winsorized Z-score (1st/99th percentile) - checks results aren't
        driven by extreme outlier countries.

HOW TO RUN:
    pip install pandas numpy statsmodels linearmodels python-docx scipy
    python 02_generate_results_document.py

INPUT:
    data/global_banking_panel.csv   (produced by 01_fetch_data.py)

OUTPUT:
    outputs/Global_Banking_Stability_Results.docx
"""

import os
import numpy as np
import pandas as pd
from scipy import stats
import statsmodels.api as sm
from linearmodels.panel import PanelOLS, RandomEffects

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# CONFIGURATION
# ---------------------------------------------------------------------------

INPUT_FILE = "data/global_banking_panel.csv"
OUTPUT_DIR = "outputs"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Global_Banking_Stability_Results.docx")

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x55, 0x55, 0x55)

CORE_VARS = [
    "bank_zscore", "capital_to_assets", "npl_ratio",
    "roa", "liquid_to_deposits", "gdp_growth", "inflation",
]

CRISIS_YEARS = {"GFC (2008)": 2008, "COVID (2020)": 2020}
PLACEBO_YEAR = 2004  # a year with no known systemic banking shock


# ---------------------------------------------------------------------------
# DOCX HELPERS
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GREY
    return p


def add_table(doc, dataframe, float_cols=None, pct_format=False):
    """Writes a pandas DataFrame into the Word doc as a formatted table."""
    float_cols = float_cols or []
    n_rows, n_cols = dataframe.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, col in enumerate(dataframe.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        set_cell_shading(cell, "1F3864")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)

    # Data rows
    for i in range(n_rows):
        for j, col in enumerate(dataframe.columns):
            val = dataframe.iloc[i, j]
            if col in float_cols and pd.notna(val):
                text = f"{val:.3f}"
            elif pd.isna(val):
                text = "-"
            else:
                text = str(val)
            cell = table.rows[i + 1].cells[j]
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9.5)
    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# ECONOMETRIC STEPS
# ---------------------------------------------------------------------------

def load_panel():
    df = pd.read_csv(INPUT_FILE)
    df["treatment_group"] = np.where(
        df["income_group"] == "High income", "High income", "Emerging market"
    )
    return df


def table_descriptives(df):
    rows = []
    for grp, sub in df.groupby("treatment_group"):
        for var in CORE_VARS:
            s = sub[var].dropna()
            rows.append({
                "Group": grp, "Variable": var, "N": len(s),
                "Mean": s.mean(), "Std Dev": s.std(),
                "Min": s.min(), "Max": s.max(),
            })
    out = pd.DataFrame(rows)
    return out


def chow_test(series_by_year, break_year):
    """
    Manual Chow test for a break in a single time series at `break_year`.
    Fits: y = a + b*year on (1) full sample, (2) pre-break, (3) post-break,
    then compares residual sum of squares (RSS) via an F-test.
    """
    s = series_by_year.dropna()
    years = s.index.values.astype(float)
    y = s.values.astype(float)

    def rss_of(y_sub, x_sub):
        X = sm.add_constant(x_sub)
        model = sm.OLS(y_sub, X).fit()
        return model.ssr, len(y_sub)

    rss_pool, n_pool = rss_of(y, years)

    pre_mask = years < break_year
    post_mask = years >= break_year
    if pre_mask.sum() < 3 or post_mask.sum() < 3:
        return {"break_year": break_year, "F_stat": np.nan, "p_value": np.nan,
                "n_pre": int(pre_mask.sum()), "n_post": int(post_mask.sum())}

    rss_pre, n_pre = rss_of(y[pre_mask], years[pre_mask])
    rss_post, n_post = rss_of(y[post_mask], years[post_mask])

    k = 2  # params per regression (intercept + slope)
    rss_ur = rss_pre + rss_post
    f_stat = ((rss_pool - rss_ur) / k) / (rss_ur / (n_pool - 2 * k))
    p_value = 1 - stats.f.cdf(f_stat, k, n_pool - 2 * k)

    return {"break_year": break_year, "F_stat": f_stat, "p_value": p_value,
            "n_pre": n_pre, "n_post": n_post}


def table_structural_breaks(df):
    world_avg = df.groupby("year")["bank_zscore"].mean()
    rows = []
    for label, yr in CRISIS_YEARS.items():
        res = chow_test(world_avg, yr)
        rows.append({
            "Crisis": label, "Break tested at": yr,
            "F-statistic": res["F_stat"], "P-value": res["p_value"],
            "N (pre)": res["n_pre"], "N (post)": res["n_post"],
            "Significant break (p<0.05)": "Yes" if res["p_value"] < 0.05 else "No",
        })
    return pd.DataFrame(rows)


def prep_panel_index(df):
    p = df.set_index(["country_name", "year"]).copy()
    return p


def fe_re_hausman(df):
    p = prep_panel_index(df.dropna(subset=CORE_VARS))
    y = p["bank_zscore"]
    X = sm.add_constant(p[["capital_to_assets", "npl_ratio", "roa",
                            "liquid_to_deposits", "gdp_growth", "inflation"]])

    fe_model = PanelOLS(y, X, entity_effects=True, time_effects=True)
    fe_res = fe_model.fit(cov_type="clustered", cluster_entity=True)

    re_model = RandomEffects(y, X)
    re_res = re_model.fit(cov_type="clustered", cluster_entity=True)

    # Hausman test: compare FE vs RE coefficients (excluding constant)
    common = [v for v in fe_res.params.index if v in re_res.params.index and v != "const"]
    b_fe = fe_res.params[common].values
    b_re = re_res.params[common].values
    v_fe = fe_res.cov.loc[common, common].values
    v_re = re_res.cov.loc[common, common].values
    diff = b_fe - b_re
    var_diff = v_fe - v_re
    try:
        stat = diff @ np.linalg.inv(var_diff) @ diff
        if stat < 0:
            # Can happen in finite samples when the variance-difference
            # matrix isn't positive definite - the test is not reliable
            # in this case rather than indicating a real result.
            stat, p_val = np.nan, np.nan
        else:
            p_val = 1 - stats.chi2.cdf(stat, df=len(common))
    except np.linalg.LinAlgError:
        stat, p_val = np.nan, np.nan

    reg_table = pd.DataFrame({
        "Variable": ["const"] + common,
        "FE coef.": [fe_res.params[v] for v in ["const"] + common],
        "FE p-value": [fe_res.pvalues[v] for v in ["const"] + common],
        "RE coef.": [re_res.params[v] for v in ["const"] + common],
        "RE p-value": [re_res.pvalues[v] for v in ["const"] + common],
    })

    if np.isnan(p_val):
        preferred = "Inconclusive (test statistic invalid - see note below)"
    elif p_val < 0.05:
        preferred = "Fixed Effects"
    else:
        preferred = "Random Effects"

    hausman_table = pd.DataFrame([{
        "Test": "Hausman (FE vs RE)", "Chi-sq stat": stat, "P-value": p_val,
        "Preferred model": preferred,
    }])

    return reg_table, hausman_table, fe_res, re_res


def did_table(df, crisis_year, label):
    d = df.dropna(subset=["bank_zscore"]).copy()
    d["post"] = (d["year"] >= crisis_year).astype(int)
    d["treat"] = (d["treatment_group"] == "Emerging market").astype(int)
    d["treat_post"] = d["treat"] * d["post"]

    p = d.set_index(["country_name", "year"])
    y = p["bank_zscore"]
    # With entity + time fixed effects, the standalone "treat" (time-invariant
    # per country) and "post" (common across countries each year) terms are
    # automatically absorbed by the fixed effects - only the interaction
    # term (the DiD estimate itself) remains identifiable, which is the
    # standard two-way fixed effects DiD specification.
    X = sm.add_constant(p[["treat_post"]])
    model = PanelOLS(y, X, entity_effects=True, time_effects=True)
    res = model.fit(cov_type="clustered", cluster_entity=True)

    out = pd.DataFrame({
        "Term": res.params.index,
        "Coefficient": res.params.values,
        "Std. Error": res.std_errors.values,
        "P-value": res.pvalues.values,
    })
    out.insert(0, "Crisis", label)
    return out, res


def event_study_table(df, crisis_year, label, window=3):
    d = df.copy()
    d["rel_year"] = d["year"] - crisis_year
    d = d[(d["rel_year"] >= -window) & (d["rel_year"] <= window)]
    pivot = (
        d.groupby(["rel_year", "treatment_group"])["bank_zscore"]
        .mean()
        .unstack("treatment_group")
        .reset_index()
        .rename(columns={"rel_year": "Years relative to crisis"})
    )
    pivot.insert(0, "Crisis", label)
    return pivot


def robustness_table(df):
    rows = []

    # 1. Main DiD result (2008), clustered SE - restated for transparency
    main_did, _ = did_table(df, 2008, "GFC (2008)")
    interaction_row = main_did[main_did["Term"] == "treat_post"].iloc[0]
    rows.append({
        "Check": "Main DiD (clustered SE by country)",
        "Interaction coef.": interaction_row["Coefficient"],
        "P-value": interaction_row["P-value"],
    })

    # 2. Placebo test - fake crisis year with no real shock
    placebo_did, _ = did_table(df, PLACEBO_YEAR, f"Placebo ({PLACEBO_YEAR})")
    placebo_row = placebo_did[placebo_did["Term"] == "treat_post"].iloc[0]
    rows.append({
        "Check": f"Placebo test (fake break at {PLACEBO_YEAR})",
        "Interaction coef.": placebo_row["Coefficient"],
        "P-value": placebo_row["P-value"],
    })

    # 3. Winsorized Z-score (1st/99th percentile) re-run of main DiD
    d = df.copy()
    lo, hi = d["bank_zscore"].quantile([0.01, 0.99])
    d["bank_zscore"] = d["bank_zscore"].clip(lo, hi)
    win_did, _ = did_table(d, 2008, "GFC (2008), winsorized")
    win_row = win_did[win_did["Term"] == "treat_post"].iloc[0]
    rows.append({
        "Check": "Main DiD, Z-score winsorized at 1%/99%",
        "Interaction coef.": win_row["Coefficient"],
        "P-value": win_row["P-value"],
    })

    return pd.DataFrame(rows)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    df = load_panel()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading(
        "Global Banking Sector Stability Around Systemic Crises (2000-2023)",
        level=0,
    )
    for run in title.runs:
        run.font.color.rgb = NAVY
    add_note(doc, "Raw empirical results document. All tables generated directly "
                   "from World Bank GFDD/WDI panel data. Figures are produced "
                   "separately in 03_generate_figures.py.")
    doc.add_paragraph()

    # Table 1
    add_heading(doc, "1. Descriptive Statistics (by Income Group)", level=1)
    desc = table_descriptives(df)
    add_table(doc, desc, float_cols=["Mean", "Std Dev", "Min", "Max"])

    # Table 2
    add_heading(doc, "2. Structural Break Test (Chow Test)", level=1)
    add_note(doc, "Tests whether the global average Bank Z-score trend breaks "
                   "significantly at each crisis year.")
    breaks = table_structural_breaks(df)
    add_table(doc, breaks, float_cols=["F-statistic", "P-value"])

    # Table 3
    add_heading(doc, "3. Panel Fixed Effects vs Random Effects + Hausman Test", level=1)
    reg_table, hausman_table, fe_res, re_res = fe_re_hausman(df)
    add_table(doc, reg_table, float_cols=["FE coef.", "FE p-value", "RE coef.", "RE p-value"])
    add_note(doc, "Note: the Hausman statistic can be invalid ('Inconclusive') in finite "
                   "samples when the FE/RE variance-difference matrix is not positive "
                   "definite - a well-known small-sample issue, not a coding error. If "
                   "this occurs, fall back to theory/robustness checks to choose FE vs RE.")
    add_table(doc, hausman_table, float_cols=["Chi-sq stat", "P-value"])

    # Table 4 & 5
    add_heading(doc, "4. Difference-in-Differences: 2008 Global Financial Crisis", level=1)
    add_note(doc, "Treatment = Emerging market economies. Control = High income economies. "
                   "Two-way fixed effects (country + year); standalone treat/post terms are "
                   "absorbed by the fixed effects, so only the treat x post interaction "
                   "(the DiD estimate) is reported.")
    did_2008, _ = did_table(df, 2008, "GFC (2008)")
    add_table(doc, did_2008, float_cols=["Coefficient", "Std. Error", "P-value"])

    add_heading(doc, "5. Difference-in-Differences: 2020 COVID-19 Shock", level=1)
    did_2020, _ = did_table(df, 2020, "COVID (2020)")
    add_table(doc, did_2020, float_cols=["Coefficient", "Std. Error", "P-value"])

    # Table 6
    add_heading(doc, "6. Event Study: Z-score Relative to Crisis Year", level=1)
    ev_2008 = event_study_table(df, 2008, "GFC (2008)")
    ev_2020 = event_study_table(df, 2020, "COVID (2020)")
    add_table(doc, ev_2008, float_cols=["High income", "Emerging market"])
    add_table(doc, ev_2020, float_cols=["High income", "Emerging market"])

    # Table 7
    add_heading(doc, "7. Robustness Checks", level=1)
    rob = robustness_table(df)
    add_table(doc, rob, float_cols=["Interaction coef.", "P-value"])

    doc.save(OUTPUT_FILE)
    print(f"Saved results document to: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
